from llama_index.llms.openai_like import OpenAILike
from typing import Optional, List, Dict
import json
import httpx
import time
import subprocess
import os
import fitz
import faiss
from pathlib import Path
from llama_index.core import VectorStoreIndex, Document, Settings, StorageContext, load_index_from_storage
from llama_index.vector_stores.faiss import FaissVectorStore
from llama_index.embeddings.langchain import LangchainEmbedding
from llama_index.core.node_parser import SentenceSplitter
import requests
from openai import OpenAI
import psutil
import re
from docx import Document
from llama_index.llms.anthropic import Anthropic
# import httpx
# import logging
# logging.basicConfig(level=logging.DEBUG)

class requirement_splitter:
    def __init__(self):
        self.model: Optional[str] = None
        self.key: Optional[str] = None
        self.url: Optional[str] = None

        my_client = httpx.Client(
            timeout=httpx.Timeout(60.0, read=180.0),
            limits=httpx.Limits(max_connections=10, max_keepalive_connections=5),
        )

        with open("../config.json", "r", encoding="utf-8")as config_file:
            data = json.load(config_file)
            # self.url = data.get("aliyun_url")
            self.url = data.get("claude_url")
            # self.key = data.get("aliyun_key")
            self.key = data.get("claude_key")
            # self.model = data.get("llama4_maverick_model")
            self.model = data.get("claude_model")

        self.llm = Anthropic(
            model=self.model,
            base_url=self.url,
            api_key=self.key,
            # context_window=96000,
            # is_chat_model=True,
            # is_function_calling_model=False,
            timeout=120,
            # client=my_client,
        )
    
    def build_tree(self, folder:str):
        abs_folder = os.path.abspath(f"./{folder}")
        if not os.path.isdir(abs_folder):
            raise NotADirectoryError(f"{folder} 不是有效目录")
        docx_files = [f for f in os.listdir(abs_folder) if f.lower().endswith(".docx")]

        if len(docx_files) == 0:
            raise FileNotFoundError(f"目录 {folder} 中没有找到 .docx 文件")
        elif len(docx_files) > 1:
            raise ValueError(f"目录 {folder} 中存在多个 .docx 文件，无法确定唯一文件")
        file_path = os.path.join(abs_folder, docx_files[0])
        doc = Document(file_path)
        # doc = Document(file)
        full_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
        requirement = "\n".join(full_text)

        analyse_prompt = f"""
<background>
你是一名精通AADL语言的软件工程师，正在根据一份需求文档分析其对应AADL项目的结构
</background>

<adaptive_thinking_framework>
你应当严格遵循以下思考步骤进行分析：
  - 整体概述：明确文档对应的项目
  - 章节排布：明确各章节之间的关系，（例如：按照功能分章、按照组件类型分章、或其他）
  - 组件分析：详细理解每章节内容，从aadl语言的13类组件（system、process、thread、subprogram、data、processor、memory、bus等）的角度解释文档内容（注意：有些组件并没有显示的定义在文章中，例如一些subprogram功能）
  - 构建组件树：根据上一步骤的组件内容，综合文章信息，给出树形结构的组件层次关系
</adaptive_thinking_framework>

<output_example>
以下是一个组件树的结构样例：
System: 飞机高度控制系统 (Height_Control_System)
├── Device: 气压传感器 (Pressure_Sensor)
├── Device: 空速控制指令传感器 (Control_Command_Sensor)
├── Processor: 高度计算分区 (Height_Computation_Processor)
├── Process: 高度计算任务 (Height_Calculation_Process)
│   └── Thread: 高度计算线程 (Height_Calculation_Thread)
│       └── Subprogram: 高度计算程序1 (Altitude_Calc_Program)
├── Processor: 高度响应分区 (Height_Response_Processor)
├── Process: 高度响应任务 (Height_Response_Process)
│   └── Thread: 高度响应线程 (Height_Response_Thread)
├── Device: 显示器 (Display_Device)
├── Device: 油门控制单元 (Throttle_Controller)
├── Device: 升降舵控制单元 (Elevator_Controller)
├── Device: 副翼控制单元 (Aileron_Controller)
├── Bus: PCIe总线 (PCIe_Bus_1, PCIe_Bus_2, PCIe_Bus_3)
</output_example>

<instruction>
你应当遵从以下指令，以优化结果：
- 构成组件树后，再次检查需求文件，*必须*确保没有遗漏组件，不需要给出分析过程
- *禁止*自由创造不属于需求文档的内容，严格遵循说明
- 组件关系应当遵从aadl规范（例如sunprogram应当位于thread内而不是其他位置）
</instruction>

<input>
以下是你将阅读的需求文档：
{requirement}
</input>

<output>
你将按照以下要求输出：
- 输出一个完整的组件树结构
- *禁止*给出任何多余的说明，我将进行正则提取
- 必须使用组件英文名称
<output>
"""
        response = self.llm.complete(analyse_prompt)
        with open(f"./{folder}/summary.txt", "w", encoding="utf-8") as summary_f:
            summary_f.write(str(response))
        return str(response)
    
    def split_requirement(self, folder:str):
        abs_folder = os.path.abspath(f"./{folder}")
        docx_files = [f for f in os.listdir(abs_folder) if f.lower().endswith(".docx")]
        file_path = os.path.join(abs_folder, docx_files[0])
        with open(f"./{folder}/summary.txt", "r", encoding="utf-8") as f:
            summary = f.read()
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
        requirement = "\n".join(full_text)
        prompt = f"""
<background>
你是一名精通AADL语言的软件工程师，正在根据一份需求文档和一份组件结构树，分析其对应AADL项目的组件细节
</background>

<adaptive_thinking_framework>
你应当严格遵循以下思考步骤进行分析：
  - 整体概述：明确文档对应的项目
  - 章节排布：明确各章节之间的关系，（例如：按照功能分章、按照组件类型分章、或其他）
  - 确认对应关系：根据组件树中的每个组件，确定其在原文章对应的段落（也许不止一处）
  - 分析组件属性：根据上一步骤确定的一系列段落内容，记录每个组件的持有的属性信息
  - 分析组件连接关系：分析组件间的连接关系（绑定、包含、连接等）
  - 模块组织：通过若干功能模块组织上述信息
</adaptive_thinking_framework>

<instruction>
你应当遵从以下指令，以优化结果：
  - 对于每个组件都应当详细检查原需求文件内容，*必须*确保属性没有遗漏
  - *禁止*自由发挥，构建不存在的属性
  - 连接关系*严格遵守*AS5506C语法规范，不能构建不合法的连接方式（例如process与processor应该在system绑定，而不是在其他环节）
</instruction>

<output_example>
以下给出了一个模块的输出样例：
$module start$
飞行控制模块
模块名称：
    flight_control
所含组件：
    Device: Threttle_Controller 
    Device: Elevator_Controller 
    Device: Aileron_Controller
内部子组件：
    无
功能说明：
    接受来自“高度相应模块”的Helight_Response_Processor的信号
具体行为：
    控制发动机推力（油门）
    调整飞机仰角（升降舵）
    稳定飞行方向（尾翼）
关联组件：
    processor: height_Response_Processor
    Bus: PCIe_Bus_3
$module end$
</output_example>

<input>
以下是你将阅读的文档：
【原始需求】：{requirement}
【组件树】：{summary}
</input>

<output>
你将按照以下要求输出：
  - 按照你划分的功能模块，参照example格式输出
  - ‘所含组件’: 代表当前模块内高层级组件
  - ‘内部子组件’: 代表‘所含组件’之中更细节的组件划分
  - ‘功能说明’: 概括了该模块的功能
  - ‘具体行为’: 反映了模块内各个组件的属性，功能
  - ‘关联组件’: 指代并不定义在当前模块内，但与项目中当前模块内组件有关的外部组件
  - 每个模块使用‘$module start$’与‘$module end$’分隔
  - *禁止*输出其他说明与无关内容，便于我进行正则
  - *严格遵守*我给出的结构与分隔符
  - 必须使用各个组件的英文名称
</output>
"""
        response = self.llm.complete(prompt)
        with open(f"./{folder}/output.txt", "w", encoding="utf-8") as out_f:
            out_f.write(str(response))
        return str(response)
        

def spliter(file_name):
    my_spliter = requirement_splitter()
    my_spliter.build_tree(file_name)
    print("已构建组件树")
    my_spliter.split_requirement(file_name)
    print("已分解需求")


if __name__ == "__main__":
    my_spliter = requirement_splitter()
    # target_path = "Air_conditioner"
    # target_path = "Ardupilot"
    # target_path = "Car"
    # target_path = "Control_sys"
    # target_path = "End_to_end"
    # target_path = "LAMP"
    # target_path = "Pacemaker"
    # target_path = "Periodic_task"
    # target_path = "Radar"
    # target_path = "Redundancy"
    # target_path = "ROSACE"
    # target_path = "Satellite"
    # target_path = "Solar_car"
    # target_path = "Time_trigger"
    target_path = "Vxworks"
    my_spliter.build_tree(target_path)
    print("已构建组件树")
    my_spliter.split_requirement(target_path)
    print("完成")
